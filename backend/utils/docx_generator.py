from docx import Document

def generate_docx(content, output_path):
    doc = Document()

    doc.add_heading(content["name"], 0)
    doc.add_paragraph(content["contact"])

    doc.add_heading("Summary", 1)
    doc.add_paragraph(content["summary"])

    doc.add_heading("Skills", 1)
    for s in content["skills"]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Experience", 1)
    for exp in content["experience"]:
        doc.add_paragraph(exp["role_institution"], style="List Bullet")
        for b in exp["bullets"]:
            doc.add_paragraph(b)

    doc.add_heading("Education", 1)
    for e in content["education"]:
        doc.add_paragraph(e)

    doc.save(output_path)
    