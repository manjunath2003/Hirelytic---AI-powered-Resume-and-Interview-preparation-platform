from docx.shared import Pt, RGBColor, Inches

ROLE_TITLE = "Lead Educator"

def generate_layout(doc, content):
    accent_navy = RGBColor(26, 54, 104)

    # NAME
    name_p = doc.add_paragraph()
    run = name_p.add_run(content.get("name", ""))
    run.font.size = Pt(22)
    run.font.color.rgb = accent_navy
    run.bold = True
    run.font.name = 'Helvetica'

    # CONTACT
    contact_p = doc.add_paragraph(content.get("contact", ""))
    contact_p.paragraph_format.space_after = Pt(10)

    def add_modern_header(title):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        run = h.add_run(title.upper())
        run.bold = True
        run.font.color.rgb = accent_navy
        run.font.size = Pt(10)

    # SECTIONS
    add_modern_header("Teacher Profile")
    doc.add_paragraph(content.get("summary", ""))

    add_modern_header("Experience")
    for exp in content.get("experience", []):
        p = doc.add_paragraph()
        p.add_run(exp.get("role_institution", "")).bold = True
        p.add_run(f" | {exp.get('duration', '')}").italic = True
        for bullet in exp.get("bullets", []):
            doc.add_paragraph(bullet, style='List Bullet')

    add_modern_header("Technical & Soft Skills")
    doc.add_paragraph(" | ".join(content.get("skills", [])))

    add_modern_header("Education & Professional Development")
    for edu in content.get("education", []):
        doc.add_paragraph(edu)
    for cert in content.get("certifications", []):
        doc.add_paragraph(cert)
        