from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROLE_TITLE = "TECH ENTHUSIAST | BCA GRADUATE"

def generate_layout(doc, content):
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # ---------------- HEADER ----------------
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(content.get("name", ""))
    run.bold = True
    run.font.size = Pt(20)

    role_p = doc.add_paragraph()
    role_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role_run = role_p.add_run(ROLE_TITLE)
    role_run.font.size = Pt(10)

    doc.add_paragraph("")  # spacing

    # ---------------- MAIN TABLE ----------------
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # 🔥 FIXED COLUMN WIDTHS (KEY PART)
    table.columns[0].width = Inches(2.2)   # LEFT (30%)
    table.columns[1].width = Inches(4.8)   # RIGHT (70%)

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]

    # ---------------- HEADER WITH LINE FUNCTION ----------------
    def header_with_line(cell, title, uppercase=False):
        row = cell.add_table(rows=1, cols=2).rows[0]

        # TITLE
        p1 = row.cells[0].paragraphs[0]
        run1 = p1.add_run(title.upper() if uppercase else title)
        run1.bold = True
        run1.font.size = Pt(11)

        # LINE (FULL WIDTH LOOK)
        p2 = row.cells[1].paragraphs[0]
        run2 = p2.add_run("──────────────────────────────")
        run2.font.size = Pt(10)

        cell.add_paragraph("")  # spacing

    # ---------------- LEFT COLUMN ----------------

    header_with_line(left, "Contact")
    left.add_paragraph(content.get("contact", ""))

    header_with_line(left, "Education")
    for edu in content.get("education", []):
        left.add_paragraph(edu)

    header_with_line(left, "Skills")
    for skill in content.get("skills", []):
        left.add_paragraph(skill, style="List Bullet")

    if content.get("languages"):
        header_with_line(left, "Language")
        for lang in content["languages"]:
            left.add_paragraph(lang, style="List Bullet")

    # ---------------- RIGHT COLUMN ----------------

    header_with_line(right, "Overview", True)
    if content.get("summary"):
        right.add_paragraph(content["summary"])

    header_with_line(right, "Projects", True)
    for proj in content.get("projects", []):
        p = right.add_paragraph()
        p.add_run(proj.get("title", "")).bold = True

        if proj.get("duration"):
            d = right.add_paragraph(proj["duration"])
            d.runs[0].italic = True

        right.add_paragraph(proj.get("description", ""))

    if content.get("experience"):
        header_with_line(right, "Experience", True)
        for exp in content["experience"]:
            p = right.add_paragraph()
            p.add_run(exp.get("role_institution", "")).bold = True

            for bullet in exp.get("bullets", []):
                right.add_paragraph(bullet, style="List Bullet")

    if content.get("certifications"):
        header_with_line(right, "Certifications", True)
        for cert in content["certifications"]:
            right.add_paragraph(cert, style="List Bullet")
            