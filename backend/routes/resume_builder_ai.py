# backend/routes/resume_builder_ai.py
"""
AI Resume Generation endpoints.

Endpoints:
- POST  /api/resume_ai/generate_ai/<user_id>
- POST  /api/resume_ai/generate/<user_id>
- GET   /api/resume_ai/list/<user_id>
- GET   /api/resume_ai/download/<user_id>/pdf/<filename>
- GET   /api/resume_ai/download/<user_id>/docx/<filename>
- GET   /api/resume_ai/preview/<user_id>/pdf/<filename>
"""

from flask import Blueprint, jsonify, send_file
from backend.extensions import mongo

from datetime import datetime, timezone
import os
import logging
import shutil
import subprocess
import uuid

from docx import Document

# --------------------------------------------------
# TEMPLATE IMPORT
# --------------------------------------------------
from backend.resume_templates.teacher_resume_v1 import (
    ROLE_TITLE,
    SUMMARY_TEMPLATE,
    CORE_SKILLS,
    DEFAULT_EXPERIENCE,
    DEFAULT_EDUCATION,
    DEFAULT_CERTIFICATIONS,
    DEFAULT_ADDITIONAL_INFO
)

logger = logging.getLogger("resume_builder_ai")
resume_ai_bp = Blueprint("resume_ai_bp", __name__)

# --------------------------------------------------
# STORAGE
# --------------------------------------------------
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
BASE_UPLOAD = os.path.join(BASE_DIR, "uploads", "generated")
os.makedirs(BASE_UPLOAD, exist_ok=True)


def _user_generated_dir(user_id: str) -> str:
    path = os.path.join(BASE_UPLOAD, str(user_id))
    os.makedirs(path, exist_ok=True)
    return path


def _safe_filename(name: str) -> str:
    return "".join(c for c in name if c.isalnum() or c in (" ", "-", "_", ".")).rstrip()


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _short_ts() -> str:
    return datetime.now().strftime("%Y%m%d%H%M%S%f")


# --------------------------------------------------
# RESUME CONTENT BUILDER (CRASH-PROOF)
# --------------------------------------------------
def assemble_resume_content(parsed: dict) -> dict:
    parsed = parsed or {}

    # ---------- PERSONAL ----------
    personal = parsed.get("personal", {}) or {}
    name = personal.get("name") or "Candidate Name"

    contact_parts = []
    for key in ("phone", "email", "linkedin"):
        val = personal.get(key)
        if isinstance(val, str) and val.strip():
            contact_parts.append(val.strip())
    contact = " | ".join(contact_parts)

    # ---------- SUMMARY ----------
    years = parsed.get("total_experience_years") or 1
    boards = parsed.get("boards") or "CBSE / ICSE / State Board"
    grades = parsed.get("grades") or "X–XII"

    summary = SUMMARY_TEMPLATE.format(
        years=years,
        boards=boards,
        grades=grades
    )

    # ---------- SKILLS ----------
    raw_skills = parsed.get("skills")
    skills = []

    if isinstance(raw_skills, list):
        for s in raw_skills:
            if isinstance(s, str) and len(s.strip()) > 2:
                skills.append(s.strip())

    if len(skills) < 5:
        skills = CORE_SKILLS

    # ---------- EXPERIENCE ----------
    raw_jobs = parsed.get("jobs")
    experience = []

    if isinstance(raw_jobs, list):
        for job in raw_jobs:
            if isinstance(job, dict):
                role = job.get("role") or job.get("title")
                bullets = job.get("bullets")

                if isinstance(role, str) and isinstance(bullets, list):
                    clean_bullets = [
                        b for b in bullets
                        if isinstance(b, str) and len(b.strip()) > 3
                    ]
                    if clean_bullets:
                        experience.append({
                            "role": role,
                            "bullets": clean_bullets
                        })

    if not experience:
        experience = DEFAULT_EXPERIENCE

    # ---------- EDUCATION ----------
    education = parsed.get("education")
    if not isinstance(education, list) or not education:
        education = DEFAULT_EDUCATION

    # ---------- CERTIFICATIONS ----------
    certifications = parsed.get("certifications")
    if not isinstance(certifications, list) or not certifications:
        certifications = DEFAULT_CERTIFICATIONS

    return {
        "name": name,
        "title": ROLE_TITLE,
        "contact": contact,
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "education": education,
        "certifications": certifications,
        "additional": DEFAULT_ADDITIONAL_INFO,
    }


# --------------------------------------------------
# DOCX GENERATOR
# --------------------------------------------------
def create_docx(content: dict, out_path: str) -> None:
    doc = Document()

    doc.add_heading(content["name"], level=1)
    doc.add_paragraph(content["title"])
    doc.add_paragraph(content["contact"])

    def section(title: str):
        p = doc.add_paragraph()
        p.add_run(title).bold = True

    section("PROFESSIONAL SUMMARY")
    doc.add_paragraph(content["summary"])

    section("CORE SKILLS")
    for skill in content["skills"]:
        doc.add_paragraph(skill, style="List Bullet")

    section("PROFESSIONAL EXPERIENCE")
    for job in content["experience"]:
        doc.add_paragraph(job["role"], style="List Number")
        for bullet in job["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet 2")

    section("EDUCATION")
    for edu in content["education"]:
        doc.add_paragraph(str(edu))

    section("CERTIFICATIONS")
    for cert in content["certifications"]:
        doc.add_paragraph(cert, style="List Bullet")

    doc.save(out_path)


# --------------------------------------------------
# DOCX → PDF (ONLY SOURCE OF TRUTH)
# --------------------------------------------------
def create_pdf_from_docx(docx_path: str, pdf_path: str) -> bool:
    libre = shutil.which("libreoffice") or shutil.which("soffice")
    if not libre:
        logger.error("LibreOffice not found in PATH")
        return False

    try:
        subprocess.run(
            [
                libre,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                os.path.dirname(pdf_path),
                docx_path,
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        generated = docx_path.replace(".docx", ".pdf")
        if os.path.exists(generated):
            shutil.move(generated, pdf_path)
            return True
    except Exception as e:
        logger.exception(e)

    return False


# --------------------------------------------------
# ROUTES
# --------------------------------------------------
@resume_ai_bp.route("/generate_ai/<user_id>", methods=["POST"])
@resume_ai_bp.route("/generate/<user_id>", methods=["POST"])
def generate_ai_resume(user_id):
    try:
        print("🔥 GENERATE RESUME ROUTE HIT 🔥", user_id)

        # ----------------------------
        # FETCH PARSED RESUME
        # ----------------------------
        parsed = mongo.db.resumes.find_one({"user_id": user_id}) or {}

        # ----------------------------
        # BUILD CONTENT
        # ----------------------------
        content = assemble_resume_content(parsed)

        # ----------------------------
        # FILENAMES
        # ----------------------------
        resume_id = str(uuid.uuid4())
        base = f"resume_{resume_id}"

        docx_filename = f"{base}.docx"
        pdf_filename = f"{base}.pdf"

        output_dir = _user_generated_dir(user_id)
        docx_path = os.path.join(output_dir, docx_filename)
        pdf_path = os.path.join(output_dir, pdf_filename)

        # ----------------------------
        # CREATE DOCX
        # ----------------------------
        create_docx(content, docx_path)

        # ----------------------------
        # CREATE PDF (OPTIONAL)
        # ----------------------------
        pdf_ok = create_pdf_from_docx(docx_path, pdf_path)

        # ----------------------------
        # SAVE METADATA
        # ----------------------------
        mongo.db.generated_resumes.insert_one({
            "user_id": user_id,
            "docx_filename": docx_filename,
            "pdf_filename": pdf_filename if pdf_ok else None,
            "created_at": datetime.utcnow()
        })

        return jsonify({
            "message": "Resume generated successfully",
            "docx_filename": docx_filename,
            "pdf_filename": pdf_filename if pdf_ok else None
        }), 200

    except Exception as e:
        print("❌ RESUME GENERATION ERROR:", e)
        logging.exception("Resume generation failed")
        return jsonify({"error": str(e)}), 500


@resume_ai_bp.route("/list/<user_id>", methods=["GET"])
def list_generated(user_id):
    items = mongo.db.generated_resumes.find({"user_id": user_id}).sort("created_at", -1)
    return jsonify([{**i, "_id": str(i["_id"])} for i in items])


@resume_ai_bp.route("/download/<user_id>/<fmt>/<filename>", methods=["GET"])
def download_file(user_id, fmt, filename):
    path = os.path.join(_user_generated_dir(user_id), filename)
    if not os.path.exists(path):
        return jsonify({"error": "File not found"}), 404
    
    mongo.db.resume_downloads.insert_one({
    "user_id": user_id,            # STRING
    "resume_id": filename,         # or resume_id if you have it
    "downloaded_at": datetime.utcnow()
})

    return send_file(path, as_attachment=True)


@resume_ai_bp.route("/preview/<user_id>/pdf/<filename>", methods=["GET"])
def preview_pdf(user_id, filename):
    path = os.path.join(_user_generated_dir(user_id), filename)
    if not os.path.exists(path):
        return jsonify({"error": "File not found"}), 404
    return send_file(path, as_attachment=False)

# --------------------------------------------------
# DELETE GENERATED RESUME
# --------------------------------------------------
@resume_ai_bp.route("/delete/<resume_id>", methods=["DELETE"])
def delete_resume(resume_id):
    from bson import ObjectId

    resume = mongo.db.generated_resumes.find_one(
        {"_id": ObjectId(resume_id)}
    )

    if not resume:
        return jsonify({"error": "Resume not found"}), 404

    user_id = resume["user_id"]

    # delete files
    for key in ["pdf_filename", "docx_filename"]:
        fname = resume.get(key)
        if fname:
            path = os.path.join(_user_generated_dir(user_id), fname)
            if os.path.exists(path):
                os.remove(path)

    mongo.db.generated_resumes.delete_one(
        {"_id": ObjectId(resume_id)}
    )

    return jsonify({"message": "Resume deleted"}), 200

# --------------------------------------------------
# RENAME GENERATED RESUME
# --------------------------------------------------
@resume_ai_bp.route("/rename/<resume_id>", methods=["PUT"])
def rename_generated_resume(resume_id):
    from flask import request
    from bson import ObjectId

    data = request.get_json(silent=True) or {}
    new_name = data.get("new_name", "").strip()

    if not new_name:
        return jsonify({"error": "Invalid name"}), 400

    resume = mongo.db.generated_resumes.find_one(
        {"_id": ObjectId(resume_id)}
    )

    if not resume:
        return jsonify({"error": "Resume not found"}), 404

    user_id = resume["user_id"]
    user_dir = _user_generated_dir(user_id)

    # ---------- CHECK DUPLICATE ----------
    existing = mongo.db.generated_resumes.find_one({
        "user_id": user_id,
        "$or": [
            {"pdf_filename": f"{new_name}.pdf"},
            {"docx_filename": f"{new_name}.docx"}
        ]
    })

    if existing:
        return jsonify({"error": "Duplicate name"}), 409

    # ---------- OLD FILES ----------
    old_pdf = resume.get("pdf_filename")
    old_docx = resume.get("docx_filename")

    new_pdf = f"{new_name}.pdf" if old_pdf else None
    new_docx = f"{new_name}.docx" if old_docx else None

    # ---------- RENAME FILES ----------
    if old_pdf:
        os.rename(
            os.path.join(user_dir, old_pdf),
            os.path.join(user_dir, new_pdf)
        )

    if old_docx:
        os.rename(
            os.path.join(user_dir, old_docx),
            os.path.join(user_dir, new_docx)
        )

    # ---------- UPDATE DB ----------
    mongo.db.generated_resumes.update_one(
        {"_id": ObjectId(resume_id)},
        {"$set": {
            "pdf_filename": new_pdf,
            "docx_filename": new_docx
        }}
    )

    return jsonify({
        "pdf_filename": new_pdf,
        "docx_filename": new_docx
    }), 200
