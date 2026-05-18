# backend/routes/resume_parser.py
# Full Resume Parsing + Ontology Mapping + Structured Output (No Raw Text Stored)

from flask import Blueprint, request, jsonify
from backend.extensions import mongo
from backend.ontology.teacher_ontology import map_teacher_ontology
from flask_cors import CORS, cross_origin

import os
import fitz  # PyMuPDF
import docx
import spacy
import uuid
import logging
from werkzeug.utils import secure_filename
from datetime import datetime, timezone
import re
import json
from transformers import pipeline
from sentence_transformers import SentenceTransformer, util
import numpy as np

# Fuzzy string matching
from rapidfuzz import fuzz
from pathlib import Path

# -------------------------------------------------------
# Logger setup
# -------------------------------------------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("resume_parser")

resume_parser_bp = Blueprint("resume_parser_bp", __name__)

# -------------------------------------------------------
# Disable caching for all responses in this Blueprint
# -------------------------------------------------------
@resume_parser_bp.after_request
def add_header(response):
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response


# -------------------------------------------------------
# Teacher requirements loading + category parsing + scoring
# -------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parents[1]  # backend/
TEACHER_REQ_PATH = BASE_DIR / "data" / "teacher_requirements.json"
_teacher_req_cache = None
TEACHER_REQUIREMENTS = {}


def load_teacher_requirements():
    global _teacher_req_cache, TEACHER_REQUIREMENTS
    if _teacher_req_cache is None:
        try:
            with open(TEACHER_REQ_PATH, "r", encoding="utf-8") as f:
                _teacher_req_cache = json.load(f)
                TEACHER_REQUIREMENTS = _teacher_req_cache
                logger.info("✅ Loaded teacher_requirements.json")
        except Exception as e:
            logger.warning(f"⚠️ Could not load teacher_requirements.json: {e}")
            _teacher_req_cache = {}
            TEACHER_REQUIREMENTS = {}
    return _teacher_req_cache or {}


# load once at import time so TEACHER_REQUIREMENTS is populated
load_teacher_requirements()


def parse_selected_category(selected_category: str):
    """
    Converts UI string into:
    - base_category_key
    - board_key
    - puc_stream_key
    """
    base_category_key = None
    board_key = None
    puc_stream_key = None

    if not selected_category:
        return None, None, None

    s = selected_category.strip()

    # Direct teaching levels
    if s == "Pre-Primary (Nursery-KG)":
        base_category_key = "Pre-Primary (Nursery-KG)"

    elif s == "Primary (Grades 1-5)":
        base_category_key = "Primary (Grades 1-5)"

    elif s == "Middle School (Grades 6-8)":
        base_category_key = "Middle School (Grades 6-8)"

    elif s.startswith("High School (Grades 9-10)"):
        base_category_key = "High School (Grades 9-10)"
        if "CBSE" in s:
            board_key = "CBSE"
        elif "ICSE" in s:
            board_key = "ICSE"
        elif "State" in s:
            board_key = "State Board"

    elif s.startswith("High School (Grades 9-12)"):
        base_category_key = "High School (Grades 9-12)"
        if "CBSE" in s:
            board_key = "CBSE"
        elif "ICSE" in s:
            board_key = "ICSE"
        elif "State" in s:
            board_key = "State Board"

    elif s.startswith("PUC –") or s.startswith("PUC -") or s.startswith("PUC "):
        base_category_key = "Higher Secondary / PUC (Grades 11-12)"
        if "PCMB" in s:
            puc_stream_key = "PCMB"
        elif "PCMCs" in s:
            puc_stream_key = "PCMCs"
        elif "Commerce" in s:
            puc_stream_key = "Commerce"

    return base_category_key, board_key, puc_stream_key


def score_resume_with_requirements(parsed_resume: dict, selected_category: str):
    """
    Uses teacher_requirements.json to compute:
    - strength
    - missing skills
    - recommendations
    - dataset_match_score
    """
    req = load_teacher_requirements()
    if not req:
        return 0, [], ["Requirements file missing or invalid."], 0

    base_key, board_key, puc_stream_key = parse_selected_category(selected_category)

    required_skill_pool = []

    # 1) Main teaching level
    if base_key and base_key in req:
        required_skill_pool.extend(req[base_key].get("core_skills", []))

    # 2) Board-specific (CBSE / ICSE / State)
    if board_key:
        boards = req.get("Boards", {})
        board_obj = boards.get(board_key, {})
        required_skill_pool.extend(board_obj.get("core_focus", []))

    # 3) PUC Streams (PCMB / PCMCs / Commerce)
    if puc_stream_key:
        streams = req.get("PUC Streams", {})
        stream_obj = streams.get(puc_stream_key, {})
        required_skill_pool.extend(stream_obj.get("core_skills", []))

    # 4) General teacher competencies
    general = req.get("General Teaching Competencies", {})
    required_skill_pool.extend(general.get("core_skills", []))

    # 5) NEP 2020 skills
    nep = req.get("NEP 2020 Competencies", {})
    required_skill_pool.extend(nep.get("core_focus", []))

    # Normalize required skills
    required_skills = sorted({s.strip().lower() for s in required_skill_pool if s.strip()})

    # Resume skills
    resume_skills_raw = parsed_resume.get("skills") or []
    resume_skills = [str(s).strip().lower() for s in resume_skills_raw if str(s).strip()]

    if not required_skills or not resume_skills:
        return 0, required_skills, ["Not enough skill information to compare."], 0

    matched = []
    for r in required_skills:
        for rs in resume_skills:
            if r in rs or rs in r:
                matched.append(r)
                break

    matched = sorted(set(matched))
    missing = [s for s in required_skills if s not in matched]

    strength = round(100 * len(matched) / len(required_skills)) if required_skills else 0
    dataset_match_score = strength

    recommendations = []
    if missing:
        recommendations.append(
            "Consider improving or adding these skills: " + ", ".join(missing[:10])
        )
    recommendations.append(
        f"Matched {len(matched)} out of {len(required_skills)} required competencies."
    )

    return strength, missing, recommendations, dataset_match_score


# -------------------------------------------------------
# Load Teacher Resume Dataset for Semantic Comparison
# -------------------------------------------------------
DATASET_PATH = BASE_DIR / "data" / "teacher_resumes_dataset.json"
JOB_DESC_PATH = BASE_DIR / "data" / "teacher_job_descriptions.json"

try:
    with open(DATASET_PATH, "r", encoding="utf-8") as f:
        TEACHER_RESUME_DATASET = json.load(f)
        logger.info(f"✅ Loaded {len(TEACHER_RESUME_DATASET)} synthetic teacher resumes.")
except Exception as e:
    TEACHER_RESUME_DATASET = []
    logger.warning(f"⚠️ Could not load teacher dataset: {e}")

# -------------------------------------------------------
# Load Teacher Job Descriptions
# -------------------------------------------------------
JOB_DESC_PATH = os.path.join(os.getcwd(), "backend", "data", "teacher_job_descriptions.json")

try:
    with open(JOB_DESC_PATH, "r", encoding="utf-8") as f:
        TEACHER_JOB_DESCRIPTIONS = json.load(f)
        logger.info("✅ Loaded teacher job description dataset.")
except Exception as e:
    TEACHER_JOB_DESCRIPTIONS = {}
    logger.warning(f"⚠️ Could not load job description dataset: {e}")

# -------------------------------------------------------
# Load NLP Models
# -------------------------------------------------------
try:
    nlp = spacy.load("en_core_web_trf")
    logger.info("✅ Loaded transformer-based spaCy model (en_core_web_trf)")
except Exception:
    nlp = spacy.load("en_core_web_sm")
    logger.warning("⚠️ Falling back to spaCy small model (en_core_web_sm)")

try:
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
    classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")
    logger.info("✅ Loaded Hugging Face summarization and classifier models")
except Exception as e:
    summarizer = None
    classifier = None
    logger.warning(f"⚠️ Failed to load Transformers pipelines: {e}")

# Load sentence transformer model for semantic similarity
try:
    similarity_model = SentenceTransformer("all-MiniLM-L6-v2")
    logger.info("✅ Loaded sentence-transformer model for semantic similarity")
except Exception as e:
    similarity_model = None
    logger.warning(f"⚠️ Could not load semantic similarity model: {e}")

# -------------------------------------------------------
# Optional OCR
# -------------------------------------------------------
try:
    import pytesseract
    from pdf2image import convert_from_bytes

    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

UPLOAD_FOLDER = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

ALLOWED_EXT = {".pdf", ".docx", ".txt"}

# -------------------------------------------------------
# Extractors
# -------------------------------------------------------
def extract_pdf(path_or_bytes):
    text = ""
    if isinstance(path_or_bytes, str):
        doc = fitz.open(path_or_bytes)
    else:
        doc = fitz.open(stream=path_or_bytes, filetype="pdf")

    try:
        for page in doc:
            t = page.get_text()
            if t.strip():
                text += t + "\n"
            elif OCR_AVAILABLE:
                img = page.get_pixmap().samples
                text += pytesseract.image_to_string(img) + "\n"
    finally:
        doc.close()

    return text


def extract_docx(path_or_bytes):
    if isinstance(path_or_bytes, str):
        document = docx.Document(path_or_bytes)
    else:
        import tempfile

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(path_or_bytes)
            tmp.flush()
            path = tmp.name
        document = docx.Document(path)
        os.remove(path)

    return "\n".join(p.text for p in document.paragraphs if p.text.strip())


def extract_txt(path_or_bytes):
    if isinstance(path_or_bytes, str):
        return open(path_or_bytes, "r", encoding="utf-8", errors="ignore").read()
    return path_or_bytes.decode("utf-8", errors="ignore")


# -------------------------------------------------------
# Preprocess
# -------------------------------------------------------
def preprocess_text(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"Page\s*\d+.*", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[•●▪■◆◇◦]", "-", text)
    text = re.sub(r"[ \t]+", " ", text)
    return "\n".join([ln.strip() for ln in text.split("\n") if ln.strip()])


# -------------------------------------------------------
# Section Classification
# -------------------------------------------------------
SECTION_HEADERS = {
    "education": ["education", "qualification", "degree"],
    "experience": ["experience", "employment", "work history"],
    "skills": ["skills", "technical skills"],
    "certifications": ["certifications", "licenses"],
    "personal": ["personal", "contact"],
}


def extract_section(text, keywords):
    pattern = r"(?mi)^(?:{})\s*[:\-]?\s*$".format("|".join(re.escape(k) for k in keywords))
    matches = list(re.finditer(pattern, text))
    if not matches:
        return ""
    start = matches[0].end()
    remainder = text[start:]
    parts = remainder.split("\n\n", 1)
    return parts[0].strip() if parts else remainder.strip()


def classify_sections(text):
    return {k: extract_section(text, v) for k, v in SECTION_HEADERS.items()}


# -------------------------------------------------------
# NER Extract
# -------------------------------------------------------
def ner_extract(text):
    if not text.strip():
        return []
    try:
        doc = nlp(text)
        return [{"text": ent.text, "label": ent.label_} for ent in doc.ents]
    except Exception:
        return []


# -------------------------------------------------------
# Completeness Score
# -------------------------------------------------------
FIELD_WEIGHTS = {
    "personal": 15,
    "professional_summary": 15,
    "skills": 15,
    "education": 12,
    "jobs": 15,
    "certifications": 5,
    "projects": 6,
    "achievements": 6,
    "languages": 4,
    "hobbies": 2,
    "subjects": 5,
    "grade_levels": 5,
}


def compute_completeness_score(parsed):
    score = 0
    personal = parsed.get("personal", {})
    if personal.get("name") and (personal.get("phone") or personal.get("email")):
        score += FIELD_WEIGHTS["personal"]

    for field, weight in FIELD_WEIGHTS.items():
        if field == "personal":
            continue
        if parsed.get(field):
            score += weight

    return min(score, 100)


# -------------------------------------------------------
# Parse Resume
# -------------------------------------------------------
def parse_resume_text_pipeline(text):
    pre = preprocess_text(text)
    sections = classify_sections(pre)
    ner_sec = {k: ner_extract(v) for k, v in sections.items()}
    mapped = map_teacher_ontology(ner_sec) or {}
    mapped["skills"] = mapped.get("skills", []) or sections.get("skills", "").split("\n")

    lines = [ln for ln in pre.split("\n") if ln]
    summary = mapped.get("professional_summary") or " ".join(lines[:4])

    parsed = {
        "preprocessed_preview": pre[:1500],
        "professional_summary": summary,
        "personal": {
            "name": mapped.get("name") or "",
            "phone": mapped.get("phone") or "",
            "email": mapped.get("email") or "",
            "linkedin": mapped.get("linkedin") or "",
        },
        "skills": mapped.get("skills") or [],
        "education": mapped.get("degrees") or [],
        "jobs": [],
        "certifications": mapped.get("certifications") or [],
        "projects": mapped.get("projects") or [],
        "achievements": mapped.get("achievements") or [],
        "subjects": mapped.get("subjects") or [],
        "grade_levels": mapped.get("grade_levels") or [],
        "languages": mapped.get("languages") or [],
        "hobbies": mapped.get("hobbies") or [],
        "ner_sections": ner_sec,
        "mapped": mapped,
    }

    exp = sections.get("experience", "")
    if exp:
        blocks = exp.split("\n\n")
        for blk in blocks:
            l = [x for x in blk.split("\n") if x]
            if not l:
                continue
            parsed["jobs"].append(
                {
                    "title": l[0] if len(l) > 0 else "",
                    "organization": l[1] if len(l) > 1 else "",
                    "description": " ".join(l[2:]) if len(l) > 2 else "",
                }
            )

    return parsed


def detect_teacher_type(parsed_resume, best_role=None):
    """
    Automatically detects teacher type based on:
    - Job-description match (best_role)
    - Experience job titles
    - Skills
    - Education level
    - Keyword presence in resume
    """

    exp_titles = " ".join([j.get("title", "").lower() for j in parsed_resume.get("jobs", [])])
    skill_text = " ".join([s.lower() for s in parsed_resume.get("skills", [])])
    edu_text = " ".join([e.lower() for e in parsed_resume.get("education", [])])

    resume_text = f"{exp_titles} {skill_text} {edu_text}"

    # 1️⃣ Highest Accuracy: Use JD match
    if best_role:
        if "Primary" in best_role:
            return "Primary School Teacher (Grades 1-5)"
        if "Middle" in best_role:
            return "Middle School Teacher (Grades 6-8)"
        if "High School" in best_role:
            return "High School Teacher (Grades 9-10)"
        if "Senior Secondary" in best_role or "PUC" in best_role:
            return "Senior Secondary / PUC Teacher (Grades 11-12)"

    # 2️⃣ Use keywords in Experience Titles
    if any(x in exp_titles for x in ["primary", "elementary", "grade 1", "grade 2", "grade 3", "grade 4", "grade 5"]):
        return "Primary School Teacher (Grades 1-5)"

    if any(x in exp_titles for x in ["middle school", "grade 6", "grade 7", "grade 8"]):
        return "Middle School Teacher (Grades 6-8)"

    if any(x in exp_titles for x in ["high school", "secondary teacher", "grade 9", "grade 10"]):
        return "High School Teacher (Grades 9-10)"

    if any(x in exp_titles for x in ["puc", "grade 11", "grade 12", "junior college"]):
        return "Senior Secondary / PUC Teacher (Grades 11-12)"

    # 3️⃣ Education-based detection
    if "d.el.ed" in edu_text:
        return "Primary School Teacher (Grades 1-5)"

    if "b.ed" in edu_text and any(x in edu_text for x in ["science", "math", "social", "english"]):
        return "Middle School Teacher (Grades 6-8)"

    if "m.sc" in edu_text or "m.a" in edu_text:
        return "Senior Secondary / PUC Teacher (Grades 11-12)"

    # 4️⃣ Skill-based detection
    if "phonics" in skill_text or "child-centered" in skill_text:
        return "Primary School Teacher (Grades 1-5)"

    if "board exam" in skill_text or "h.o.t.s" in skill_text:
        return "High School Teacher (Grades 9-10)"

    # 5️⃣ Final fallback
    return "High School Teacher (Grades 9-10)"


# -------------------------------------------------------
# Semantic Helper Functions
# -------------------------------------------------------
def is_semantically_similar(term, list_to_check, threshold=0.55):
    if not list_to_check:
        return False

    term_lower = term.lower().strip()

    # Normalize list
    norm_list = [x.lower().strip() for x in list_to_check]

    # ✅ 1. Exact match
    if term_lower in norm_list:
        return True

    # ✅ 2. Fuzzy semantic match (NEW)
    for item in norm_list:
        try:
            if fuzz.partial_ratio(term_lower, item) >= 80:
                return True
        except Exception:
            pass

    # Sentence transformer similarity (optional)
    if similarity_model:
        try:
            term_emb = similarity_model.encode(term_lower, convert_to_tensor=True)
            list_emb = similarity_model.encode(norm_list, convert_to_tensor=True)
            cos_scores = util.cos_sim(term_emb, list_emb)[0]
            if float(np.max(cos_scores)) >= threshold:
                return True
        except Exception:
            pass

    try:
        doc1 = nlp(term_lower)
        for s in norm_list:
            doc2 = nlp(s)
            if doc1.similarity(doc2) >= threshold:
                return True
    except Exception:
        pass

    return False


def calculate_resume_strength_semantic(
    parsed_resume,
    teacher_type="High School (Grades 9-10)",
    board="CBSE",
    stream=None,
    best_role=None,
):
    if not parsed_resume or not TEACHER_REQUIREMENTS:
        return 0, []

    if teacher_type == "High School (Grades 9-12)":
        teacher_type = "High School (Grades 9-10)"

    ref_data = TEACHER_REQUIREMENTS.get(teacher_type, {})
    board_data = TEACHER_REQUIREMENTS.get("Boards", {}).get(board, {})

    required_skills = ref_data.get("core_skills", [])
    required_qualifications = ref_data.get("required_qualifications", [])
    required_experience = ref_data.get("preferred_experience", [])
    board_focus = board_data.get("core_focus", [])

    resume_skills = parsed_resume.get("skills", [])
    resume_education = parsed_resume.get("education", [])
    resume_jobs = [j.get("title", "") for j in parsed_resume.get("jobs", [])]

    total_checks = 0
    matched_checks = 0
    missing_concepts = []

    def evaluate(term, data_list):
        nonlocal total_checks, matched_checks
        total_checks += 1

        if is_semantically_similar(term, data_list):
            matched_checks += 1
        else:
            missing_concepts.append(term)

    for q in required_qualifications:
        evaluate(q, resume_education)

    for s in required_skills:
        evaluate(s, resume_skills)

    for e in required_experience:
        evaluate(e, resume_jobs)

    for f in board_focus:
        evaluate(f, resume_skills + resume_jobs)

    jd = TEACHER_JOB_DESCRIPTIONS.get(best_role, {})

    job_required_skills = jd.get("required_skills", [])
    job_responsibilities = jd.get("responsibilities", [])
    job_required_qualifications = jd.get("required_qualifications", [])

    for concept in (job_required_skills + job_responsibilities + job_required_qualifications):
        evaluate(concept, resume_skills + resume_jobs + resume_education)

    if total_checks == 0:
        return 0, missing_concepts

    score = min(100, round((matched_checks / total_checks) * 100 / 0.9, 2))
    score = max(score, 80)
    return score, missing_concepts


def detect_skill_gaps(parsed_resume, teacher_type="High School (Grades 9-10)", board="CBSE", stream=None,):
    reference = TEACHER_REQUIREMENTS.get(teacher_type, {})
    board_ref = TEACHER_REQUIREMENTS.get("Boards", {}).get(board, {})

    required_skills = reference.get("core_skills", [])
    required_qualifications = reference.get("required_qualifications", [])
    required_experience = reference.get("preferred_experience", [])
    board_focus = board_ref.get("core_focus", [])

    resume_skills = parsed_resume.get("skills", [])
    resume_education = parsed_resume.get("education", [])
    resume_jobs = [j.get("title", "") for j in parsed_resume.get("jobs", [])]

    missing_skills = [s for s in required_skills if not is_semantically_similar(s, resume_skills)]
    missing_qualifications = [
        q for q in required_qualifications if not is_semantically_similar(q, resume_education)
    ]
    missing_experience = [e for e in required_experience if not is_semantically_similar(e, resume_jobs)]
    missing_board_focus = [b for b in board_focus if not is_semantically_similar(b, resume_skills)]

    recommendations = []
    if missing_skills:
        recommendations.append(f"Develop or highlight skills like: {', '.join(missing_skills[:5])}.")
    if missing_qualifications:
        recommendations.append(f"Consider obtaining or emphasizing: {', '.join(missing_qualifications[:3])}.")
    if missing_experience:
        recommendations.append(
            "Add more teaching or leadership experiences related to subject delivery or exam preparation."
        )
    if missing_board_focus:
        recommendations.append(f"Align more with {board} pedagogy areas such as: {', '.join(missing_board_focus[:3])}.")

    return {
        "missing_skills": missing_skills,
        "missing_qualifications": missing_qualifications,
        "missing_experience": missing_experience,
        "missing_board_focus": missing_board_focus,
        "recommendations": recommendations,
    }


def compare_with_dataset(parsed_resume):
    if not similarity_model or not TEACHER_RESUME_DATASET:
        return 0, {}

    resume_skills = parsed_resume.get("skills", [])
    resume_education = parsed_resume.get("education", [])
    resume_experience = [j.get("title", "") for j in parsed_resume.get("jobs", [])]

    resume_text = " ".join(resume_skills + resume_education + resume_experience).lower()
    query_emb = similarity_model.encode(resume_text, convert_to_tensor=True)

    best_match = None
    best_score = 0

    for item in TEACHER_RESUME_DATASET:
        ref_text = " ".join(
            item.get("skills", []) + item.get("education", []) + item.get("experience", [])
        ).lower()
        ref_emb = similarity_model.encode(ref_text, convert_to_tensor=True)
        score = float(util.cos_sim(query_emb, ref_emb)[0])
        if score > best_score:
            best_score = score
            best_match = item

    return round(best_score * 100, 2), best_match


def match_with_job_descriptions(parsed_resume):
    resume_text = (
        " ".join(parsed_resume.get("skills", []))
        + " "
        + " ".join(parsed_resume.get("education", []))
        + " "
        + " ".join([j.get("title", "") for j in parsed_resume.get("jobs", [])])
    ).lower()

    best_match = None
    best_score = 0

    for role, jd in TEACHER_JOB_DESCRIPTIONS.items():
        jd_text = (
            " ".join(jd.get("required_skills", []))
            + " "
            + " ".join(jd.get("responsibilities", []))
            + " "
            + " ".join(jd.get("required_qualifications", []))
        ).lower()

        score = fuzz.partial_ratio(resume_text, jd_text)

        if score > best_score:
            best_score = score
            best_match = role

    return best_match, best_score


# -------------------------------------------------------
# UPLOAD ROUTE
# -------------------------------------------------------
from flask_cors import cross_origin

@resume_parser_bp.route("/upload", methods=["POST", "OPTIONS"])
@cross_origin(
    origins="http://localhost:3000",
    supports_credentials=True,
    allow_headers=["Content-Type", "Authorization", "Accept"]
)
def upload_resume():
    if request.method == "OPTIONS":
        return jsonify({"status": "ok"}), 200

    try:
        file = request.files.get("file")
        user_id = request.form.get("user_id")
        selected_category = request.form.get("selected_category", "")

        if not file:
            return jsonify({"error": "No file uploaded"}), 400
        if not user_id:
            return jsonify({"error": "Missing user_id"}), 400
        if not selected_category:
            return jsonify({"error": "Missing teaching category"}), 400

        filename = secure_filename(file.filename)
        ext = os.path.splitext(filename)[1].lower()

        if ext not in ALLOWED_EXT:
            return jsonify({"error": "Invalid file type"}), 400

        # ------------------------------------------------
        # SAVE RAW FILE IN MONGODB FOR VIEWING LATER
        # ------------------------------------------------
        resume_uid = uuid.uuid4().hex
        file_bytes = file.read()
        file.stream.seek(0)

        mongo.db.resume_files.insert_one({
            "_id": resume_uid,      # this is the resume_id to store everywhere
            "user_id": user_id,
            "filename": filename,
            "content_type": file.content_type,
            "data": file_bytes,
            "uploaded_at": datetime.now(timezone.utc).isoformat()
        })

        # ------------------------------------------------
        # Save temporarily to parse text
        # ------------------------------------------------
        save_path = os.path.join(UPLOAD_FOLDER, f"{resume_uid}_{filename}")
        file.save(save_path)

        # Extract text for parsing
        if ext == ".pdf":
            text = extract_pdf(save_path)
        elif ext == ".docx":
            text = extract_docx(save_path)
        else:
            text = extract_txt(save_path)

        # Remove temp file
        if os.path.exists(save_path):
            os.remove(save_path)

        if not text.strip():
            return jsonify({"error": "Extraction failed"}), 400

        # ------------------------------------------------
        # Continue with your existing parsing + analysis
        # ------------------------------------------------
        parsed = parse_resume_text_pipeline(text)
        completeness = compute_completeness_score(parsed)

        teacher_type, board, stream = map_category_to_backend(selected_category)

        semantic_strength = 0
        missing_concepts = []
        missing_combined = []
        recommendations = []
        dataset_match_score = 0.0

        try:
            if teacher_type:
                semantic_strength, missing_concepts = calculate_resume_strength_semantic(
                    parsed_resume=parsed,
                    teacher_type=teacher_type,
                    board=board,
                    stream=stream
                )

                gaps = detect_skill_gaps(
                    parsed_resume=parsed,
                    teacher_type=teacher_type,
                    board=board,
                    stream=stream
                )

                missing_combined = (
                    gaps.get("missing_skills", [])
                    + gaps.get("missing_qualifications", [])
                    + gaps.get("missing_experience", [])
                )
                recommendations = gaps.get("recommendations", [])

            try:
                dataset_match_score, _best = compare_with_dataset(parsed)
            except Exception as inner_e:
                dataset_match_score = 0.0

        except Exception as inner:
            pass

        uploaded_ts = datetime.now(timezone.utc).isoformat()

        mongo.db.resumes.update_one(
            {"user_id": user_id},
            {
                "$set": {
                    "parsed": parsed,
                    "resume_strength": completeness,
                    "semantic_strength": semantic_strength,
                    "dataset_match_score": dataset_match_score,
                    "missing": missing_combined,
                    "missing_concepts": missing_concepts,
                    "recommendations": recommendations,
                    "selected_category": selected_category,
                    "uploaded_at": uploaded_ts,
                    "updated_at": uploaded_ts,
                    "resume_id": resume_uid   # <-- IMPORTANT
                }
            },
            upsert=True
        )

        return jsonify({
            "message": "Resume uploaded and analyzed successfully",
            "resume_id": resume_uid,
            "parsed_preview": parsed.get("preprocessed_preview", "")[:500],
            "category_used": selected_category,
            "strength": completeness,
            "semantic_strength": semantic_strength,
            "dataset_match_score": dataset_match_score,
            "missing": missing_combined,
            "missing_concepts": missing_concepts,
            "recommendations": recommendations
        }), 200

    except Exception as e:
        logger.exception("UPLOAD ERROR")
        return jsonify({"error": str(e)}), 500


# -------------------------------------------------------
# GET LAST SAVED ANALYSIS (PERSISTED RESULTS)
# -------------------------------------------------------
@resume_parser_bp.route("/analysis/<user_id>", methods=["GET"])
def get_saved_analysis(user_id):
    rec = mongo.db.resumes.find_one({"user_id": user_id})
    if not rec:
        return jsonify({"exists": False}), 200

    return (
        jsonify(
            {
                "exists": True,
                "strength": rec.get("resume_strength", 0),
                "semantic_strength": rec.get("semantic_strength", 0),
                "dataset_match_score": rec.get("dataset_match_score", 0),
                "missing": rec.get("missing", []),
                "missing_concepts": rec.get("missing_concepts", []),
                "recommendations": rec.get("recommendations", []),
                "updated_at": rec.get("updated_at"),
                # 🔥 NEW: expose category-based analysis too
                "selected_category": rec.get("selected_category"),
                "category_strength": rec.get("category_strength", rec.get("resume_strength", 0)),
                "category_missing": rec.get("category_missing", []),
                "category_recommendations": rec.get("category_recommendations", []),
                "category_dataset_match_score": rec.get("category_dataset_match_score", rec.get("dataset_match_score", 0)),
            }
        ),
        200,
    )

def map_category_to_backend(category):
    """
    Converts frontend category string → teacher_type, board, stream
    """
    teacher_type = ""
    board = ""
    stream = ""

    # Pre Primary
    if category == "Pre-Primary (Nursery-KG)":
        teacher_type = "Pre-Primary (Nursery-KG)"
        board = "None"

    # Primary
    elif category == "Primary (Grades 1-5)":
        teacher_type = "Primary (Grades 1-5)"
        board = "None"

    # Middle School
    elif category == "Middle School (Grades 6-8)":
        teacher_type = "Middle School (Grades 6-8)"
        board = "None"

    # High School 9–10
    elif "High School (Grades 9-10)" in category:
        teacher_type = "High School (Grades 9-10)"
        if "CBSE" in category:
            board = "CBSE"
        elif "ICSE" in category:
            board = "ICSE"
        else:
            board = "State Board"

    # High School 9–12
    elif "High School (Grades 9-12)" in category:
        teacher_type = "High School (Grades 9-12)"
        if "CBSE" in category:
            board = "CBSE"
        elif "ICSE" in category:
            board = "ICSE"
        else:
            board = "State Board"

    # PUC Streams
    elif "PUC" in category:
        teacher_type = "Higher Secondary / PUC (Grades 11-12)"
        if "PCMB" in category:
            stream = "PCMB"
        elif "PCMCs" in category:
            stream = "PCMCs"
        elif "Commerce" in category:
            stream = "Commerce"

    return teacher_type, board, stream

# -------------------------------------------------------
# STRENGTH / GAP CHECK ROUTE
# -------------------------------------------------------
@resume_parser_bp.route("/strength/<user_id>", methods=["GET", "OPTIONS"])
@cross_origin(
    origins="http://localhost:3000",
    supports_credentials=True,
    allow_headers=["Content-Type", "Authorization", "Accept"]
)

def calculate_strength(user_id):
    if request.method == "OPTIONS":
        return jsonify({"status": "ok"}), 200

    try:
        rec = mongo.db.resumes.find_one({"user_id": user_id})
        if not rec:
            return jsonify({"error": "No resume found"}), 404

        parsed = rec.get("parsed", {})
        if not parsed:
            return jsonify({"error": "Parsed resume empty"}), 500

        selected_category = rec.get("selected_category", "")

        # Convert selected category
        teacher_type, board, stream = map_category_to_backend(selected_category)

        # MAIN CALCULATIONS
        semantic_strength, missing_concepts = calculate_resume_strength_semantic(
            parsed_resume=parsed,
            teacher_type=teacher_type,
            board=board,
            stream=stream
        )

        gaps = detect_skill_gaps(
            parsed_resume=parsed,
            teacher_type=teacher_type,
            board=board,
            stream=stream
        )

        missing_combined = (
            gaps.get("missing_skills", []) +
            gaps.get("missing_qualifications", []) +
            gaps.get("missing_experience", [])
        )

        recommendations = gaps.get("recommendations", [])

        updated_ts = datetime.now(timezone.utc).isoformat()

        mongo.db.resumes.update_one(
            {"user_id": user_id},
            {
                "$set": {
                    "semantic_strength": semantic_strength,
                    "missing_concepts": missing_concepts,
                    "missing": missing_combined,
                    "recommendations": recommendations,
                    "updated_at": updated_ts
                }
            }
        )

        return jsonify({
            "strength": rec.get("resume_strength", 0),
            "semantic_strength": semantic_strength,
            "missing": missing_combined,
            "missing_concepts": missing_concepts,
            "recommendations": recommendations,
            "selected_category": selected_category,
            "teacher_type_used": teacher_type,
            "board_used": board,
            "stream_used": stream,
            "updated_at": updated_ts
        }), 200

    except Exception as e:
        logger.exception("STRENGTH ERROR")
        return jsonify({"error": str(e)}), 500

from flask import send_file
import io

@resume_parser_bp.route("/get/<resume_id>", methods=["GET"])
def get_uploaded_resume(resume_id):
    try:
        file = mongo.db.resume_files.find_one({"_id": resume_id})
        if not file:
            return jsonify({"error": "Resume not found"}), 404

        return send_file(
            io.BytesIO(file["data"]),
            mimetype=file.get("content_type", "application/octet-stream"),
            download_name=file.get("filename", "resume.pdf")
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# -------------------------------------------------------
# LIST UPLOADED RESUMES FOR A USER
# -------------------------------------------------------
@resume_parser_bp.route("/resume_files/list/<user_id>", methods=["GET"])
def list_uploaded_resumes(user_id):
    try:
        files = list(mongo.db.resume_files.find({"user_id": user_id}))
        out = []

        for f in files:
            out.append({
                "_id": f["_id"],
                "filename": f.get("filename"),
                "content_type": f.get("content_type"),
                "uploaded_at": f.get("uploaded_at")
            })

        return jsonify(out), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@resume_parser_bp.route("/resume_files/view/<resume_id>", methods=["GET"])
def view_uploaded_resume(resume_id):
    try:
        file = mongo.db.resume_files.find_one({"_id": resume_id})
        if not file:
            return jsonify({"error": "File not found"}), 404

        return send_file(
            io.BytesIO(file["data"]),
            mimetype=file.get("content_type", "application/octet-stream"),
            download_name=file.get("filename")
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500
